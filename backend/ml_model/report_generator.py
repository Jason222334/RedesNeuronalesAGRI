"""
Generador de Reportes — Sistema Inteligente Agrícola
Genera reportes en PDF, Excel y Word con tablas, figuras e interpretaciones.
"""
import json
import os
import pandas as pd
import numpy as np
from pathlib import Path
from datetime import datetime
import io

OUTPUT_DIR = Path(__file__).parent / 'outputs'
FIGURES_DIR = OUTPUT_DIR / 'figures'

TITULO_PROYECTO = "Sistema Inteligente de Predicción de Precios Agrícolas"
SUBTITULO       = "Valle Jequetepeque — Análisis Comparativo de Modelos IA"
INSTITUCION     = "Universidad Nacional de Trujillo (UNT)"
FECHA           = datetime.now().strftime('%d de %B de %Y')


def cargar_resultados():
    """Carga los JSONs de resultados generados por los módulos de análisis."""
    resultados = {}
    archivos = {
        'eda':       OUTPUT_DIR / 'eda_resultados.json',
        'modelos':   OUTPUT_DIR / 'resultados_entrenamiento.json',
        'cv':        OUTPUT_DIR / 'cross_validation_resultados.json',
        'hiperpar':  OUTPUT_DIR / 'hyperparameter_resultados.json',
        'estadist':  OUTPUT_DIR / 'pruebas_estadisticas.json',
    }
    for clave, ruta in archivos.items():
        if ruta.exists():
            with open(ruta, 'r', encoding='utf-8') as f:
                resultados[clave] = json.load(f)
        else:
            resultados[clave] = None
    return resultados


def generar_pdf(output_path=None):
    """
    Genera reporte PDF completo con tablas, figuras e interpretaciones.
    Requiere: reportlab
    """
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle, Image, PageBreak,
                                     HRFlowable)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

    if output_path is None:
        output_path = OUTPUT_DIR / 'reporte_ia_agricola.pdf'
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    resultados = cargar_resultados()

    # ── Estilos
    styles = getSampleStyleSheet()
    verde  = colors.HexColor('#1a7a3e')
    verde_claro = colors.HexColor('#d5f5e3')
    gris_oscuro = colors.HexColor('#2c3e50')

    estilo_titulo = ParagraphStyle('Titulo', parent=styles['Title'],
                                    fontSize=20, textColor=verde, spaceAfter=6,
                                    alignment=TA_CENTER, fontName='Helvetica-Bold')
    estilo_subtitulo = ParagraphStyle('Subtitulo', parent=styles['Normal'],
                                       fontSize=12, textColor=gris_oscuro,
                                       spaceAfter=4, alignment=TA_CENTER)
    estilo_h1 = ParagraphStyle('H1', parent=styles['Heading1'],
                                fontSize=14, textColor=verde, fontName='Helvetica-Bold',
                                spaceBefore=12, spaceAfter=6)
    estilo_h2 = ParagraphStyle('H2', parent=styles['Heading2'],
                                fontSize=12, textColor=gris_oscuro, fontName='Helvetica-Bold',
                                spaceBefore=8, spaceAfter=4)
    estilo_body = ParagraphStyle('Body', parent=styles['Normal'],
                                  fontSize=9, leading=13, alignment=TA_JUSTIFY,
                                  spaceAfter=6)
    estilo_nota = ParagraphStyle('Nota', parent=styles['Normal'],
                                  fontSize=8, textColor=colors.grey, fontStyle='italic',
                                  spaceAfter=4)

    doc = SimpleDocTemplate(str(output_path), pagesize=A4,
                             leftMargin=2*cm, rightMargin=2*cm,
                             topMargin=2.5*cm, bottomMargin=2*cm)
    contenido = []

    def agregar_figura(ruta, ancho=15*cm, alto=9*cm, caption=''):
        if ruta and Path(ruta).exists():
            contenido.append(Image(str(ruta), width=ancho, height=alto))
            if caption:
                contenido.append(Paragraph(f'<i>Figura: {caption}</i>', estilo_nota))
            contenido.append(Spacer(1, 0.3*cm))

    def tabla_pdf(data, cabecera=None, col_widths=None):
        if cabecera:
            table_data = [cabecera] + data
        else:
            table_data = data
        t = Table(table_data, colWidths=col_widths)
        t.setStyle(TableStyle([
            ('BACKGROUND',  (0, 0), (-1, 0), verde),
            ('TEXTCOLOR',   (0, 0), (-1, 0), colors.white),
            ('FONTNAME',    (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE',    (0, 0), (-1, -1), 8),
            ('ALIGN',       (0, 0), (-1, -1), 'CENTER'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, verde_claro]),
            ('GRID',        (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ('TOPPADDING',  (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING',(0, 0), (-1, -1), 4),
        ]))
        contenido.append(t)
        contenido.append(Spacer(1, 0.3*cm))

    # ── PORTADA
    contenido.append(Spacer(1, 3*cm))
    contenido.append(Paragraph(TITULO_PROYECTO, estilo_titulo))
    contenido.append(Paragraph(SUBTITULO, estilo_subtitulo))
    contenido.append(Spacer(1, 0.5*cm))
    contenido.append(HRFlowable(width='100%', thickness=2, color=verde))
    contenido.append(Spacer(1, 0.5*cm))
    contenido.append(Paragraph(f'<b>Institución:</b> {INSTITUCION}', estilo_body))
    contenido.append(Paragraph(f'<b>Fecha de generación:</b> {FECHA}', estilo_body))
    contenido.append(Paragraph(
        '<b>Resumen:</b> Este reporte presenta el análisis comparativo de cinco modelos de '
        'inteligencia artificial para la predicción de precios agrícolas en el Valle Jequetepeque, '
        'Perú. Se incluyen tres modelos clásicos (ARIMA, Random Forest, XGBoost) y dos modelos '
        'híbridos de redes neuronales (LSTM, CNN-LSTM), evaluados mediante métricas estándar, '
        'validación cruzada y pruebas estadísticas rigurosas.', estilo_body))
    contenido.append(PageBreak())

    # ── SECCIÓN 1: EDA
    contenido.append(Paragraph('1. Análisis Exploratorio de Datos (EDA)', estilo_h1))
    if resultados.get('eda'):
        eda = resultados['eda']
        contenido.append(Paragraph(
            f"El dataset comprende <b>{eda.get('n_registros', '?')} registros</b> para el período "
            f"<b>{eda.get('periodo', '?')}</b>, cubriendo los cultivos: "
            f"{', '.join(eda.get('cultivos', []))}. Tras la limpieza se retuvieron "
            f"{eda.get('n_registros_limpios', '?')} registros.", estilo_body))

        # Tabla de estadísticos
        if eda.get('estadisticos'):
            stats_data = []
            desc = eda['estadisticos']
            for var in ['precio_s_ton', 'temperatura', 'precipitacion_mm', 'costo_transporte']:
                if var in desc:
                    d = desc[var]
                    stats_data.append([
                        var, str(d.get('count', '')), str(d.get('mean', '')),
                        str(d.get('std', '')), str(d.get('min', '')), str(d.get('max', ''))
                    ])
            tabla_pdf(stats_data,
                      cabecera=['Variable', 'N', 'Media', 'Desv. Est.', 'Mín', 'Máx'],
                      col_widths=[4*cm, 2*cm, 2.5*cm, 2.5*cm, 2*cm, 2*cm])

        agregar_figura(eda.get('figuras', {}).get('series_temporales'),
                       caption='Series temporales de precios por cultivo (con media móvil 3 meses)')
        agregar_figura(eda.get('figuras', {}).get('heatmap'),
                       caption='Mapa de calor de correlaciones entre variables del sistema')
        agregar_figura(eda.get('figuras', {}).get('distribucion'),
                       caption='Distribución de precios por cultivo')
    contenido.append(PageBreak())

    # ── SECCIÓN 2: ENTRENAMIENTO DE MODELOS
    contenido.append(Paragraph('2. Entrenamiento y Comparativa de Modelos', estilo_h1))
    contenido.append(Paragraph(
        'Se entrenaron cinco modelos: ARIMA (clásico univariado), Random Forest y XGBoost '
        '(clásicos multivariados), LSTM y CNN-LSTM (redes neuronales híbridas). '
        'La evaluación usa RMSE, MAE, R² y MAPE como métricas estándar.', estilo_body))

    if resultados.get('modelos'):
        mod = resultados['modelos']
        if mod.get('tabla_metricas'):
            tabla_data = []
            for fila in mod['tabla_metricas']:
                tabla_data.append([
                    fila.get('Modelo', ''), str(fila.get('RMSE', '')),
                    str(fila.get('MAE', '')), str(fila.get('R²', '')),
                    str(fila.get('MAPE %', '')), f"{fila.get('Tiempo s', '')}s"
                ])
            tabla_pdf(tabla_data,
                      cabecera=['Modelo', 'RMSE', 'MAE', 'R²', 'MAPE %', 'Tiempo'],
                      col_widths=[4*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2*cm])

        mejor = mod.get('mejor_modelo', 'N/D')
        r2_m  = mod.get('mejor_r2', 'N/D')
        contenido.append(Paragraph(
            f'<b>Modelo con mejor desempeño:</b> {mejor.upper()} (R² = {r2_m}). '
            f'Este modelo fue guardado como archivo .h5/.pkl para consumo en producción '
            f'sin necesidad de reentrenamiento.', estilo_body))

        agregar_figura(mod.get('figuras', {}).get('tabla_comparativa'),
                       caption='Tabla comparativa visual de modelos')
        agregar_figura(mod.get('figuras', {}).get('predicciones_vs_real'),
                       caption='Predicción vs valores reales — comparativa de los 5 modelos')
        agregar_figura(mod.get('figuras', {}).get('importancia_features'),
                       caption='Mapa de calor de importancia de variables (Random Forest)')
        agregar_figura(mod.get('figuras', {}).get('confusion_roc'),
                       caption='Matriz de confusión y curva ROC — Clasificación de riesgo de plaga')
    contenido.append(PageBreak())

    # ── SECCIÓN 3: VALIDACIÓN CRUZADA
    contenido.append(Paragraph('3. Validación Cruzada (Cross-Validation)', estilo_h1))
    if resultados.get('cv'):
        cv = resultados['cv']
        n_folds = cv.get('n_folds', 5)
        contenido.append(Paragraph(
            f'Se utilizó <b>TimeSeriesSplit con {n_folds} folds</b> (configurable), '
            f'método adecuado para datos temporales ya que evita data leakage '
            f'al respetar el orden cronológico de los datos.', estilo_body))

        tabla_cv = []
        for nombre_m in ['random_forest', 'xgboost']:
            datos_m = cv.get(nombre_m, {})
            if datos_m:
                tabla_cv.append([
                    nombre_m.replace('_', ' ').title(),
                    str(datos_m.get('RMSE_mean', '')), str(datos_m.get('RMSE_std', '')),
                    str(datos_m.get('MAE_mean', '')),  str(datos_m.get('MAE_std', '')),
                    str(datos_m.get('R2_mean', '')),   str(datos_m.get('R2_std', ''))
                ])
        if tabla_cv:
            tabla_pdf(tabla_cv,
                      cabecera=['Modelo', 'RMSE μ', 'RMSE σ', 'MAE μ', 'MAE σ', 'R² μ', 'R² σ'],
                      col_widths=[3.5*cm, 2*cm, 1.8*cm, 2*cm, 1.8*cm, 2*cm, 1.8*cm])

        agregar_figura(cv.get('figura'),
                       caption=f'Métricas de validación cruzada por fold ({n_folds} folds)')
    contenido.append(PageBreak())

    # ── SECCIÓN 4: HIPERPARÁMETROS
    contenido.append(Paragraph('4. Ajuste de Hiperparámetros (Hyperparameter Tuning)', estilo_h1))
    if resultados.get('hiperpar'):
        hp = resultados['hiperpar']
        contenido.append(Paragraph(
            'Se aplicó búsqueda bayesiana con <b>Optuna</b> para Random Forest y XGBoost, '
            'y Grid Search exhaustivo para los parámetros (p, d, q) del modelo ARIMA.', estilo_body))

        for m_name in ['random_forest', 'xgboost']:
            datos_hp = hp.get(m_name, {})
            if datos_hp and datos_hp.get('best_params'):
                contenido.append(Paragraph(f'<b>{m_name.replace("_"," ").title()}</b>:', estilo_h2))
                params_data = [[str(k), str(v)] for k, v in datos_hp['best_params'].items()]
                tabla_pdf(params_data, cabecera=['Hiperparámetro', 'Valor Óptimo'],
                          col_widths=[7*cm, 7*cm])

        if hp.get('arima') and hp['arima'].get('mejor'):
            arima_m = hp['arima']['mejor']
            contenido.append(Paragraph(
                f'<b>ARIMA:</b> Mejor orden encontrado = {arima_m.get("orden")} '
                f'(AIC = {arima_m.get("aic", ""):.2f}, RMSE = {arima_m.get("rmse", "")})',
                estilo_body))

        agregar_figura(hp.get('figura_convergencia'),
                       caption='Convergencia del proceso de optimización bayesiana (Optuna)')
    contenido.append(PageBreak())

    # ── SECCIÓN 5: PRUEBAS ESTADÍSTICAS
    contenido.append(Paragraph('5. Pruebas Estadísticas de Validación', estilo_h1))
    if resultados.get('estadist'):
        est = resultados['estadist']
        contenido.append(Paragraph(
            'Se aplicaron cinco familias de pruebas estadísticas rigurosas para validar '
            'la calidad de los modelos y comparar su desempeño de forma objetiva:', estilo_body))

        # Tabla Shapiro-Wilk
        if est.get('shapiro_wilk'):
            contenido.append(Paragraph('<b>Test de Shapiro-Wilk (Normalidad de Residuos):</b>', estilo_h2))
            sw_data = [[r['modelo'], str(r['estadistico']), str(r['p_valor']), r['resultado']]
                       for r in est['shapiro_wilk']]
            tabla_pdf(sw_data,
                      cabecera=['Modelo', 'Estadístico W', 'p-valor', 'Resultado'],
                      col_widths=[4*cm, 3.5*cm, 3*cm, 5*cm])

        # Tabla Ljung-Box
        if est.get('ljung_box'):
            contenido.append(Paragraph('<b>Test de Ljung-Box (Autocorrelación de Residuos):</b>', estilo_h2))
            lb_data = [[r['modelo'], str(r['estadistico']), str(r['p_valor']), r['resultado']]
                       for r in est['ljung_box']]
            tabla_pdf(lb_data,
                      cabecera=['Modelo', 'Estadístico LB', 'p-valor', 'Resultado'],
                      col_widths=[4*cm, 3.5*cm, 3*cm, 5*cm])

        # Tabla Diebold-Mariano
        if est.get('diebold_mariano'):
            contenido.append(Paragraph('<b>Test de Diebold-Mariano (Comparación de Modelos):</b>', estilo_h2))
            dm_data = [[r['modelos'], str(r.get('estadistico_DM', '')), str(r.get('p_valor', '')),
                        r.get('resultado', '')]
                       for r in est['diebold_mariano']]
            tabla_pdf(dm_data,
                      cabecera=['Par de Modelos', 'Estadístico DM', 'p-valor', 'Resultado'],
                      col_widths=[5.5*cm, 3*cm, 2.5*cm, 5*cm])

        agregar_figura(est.get('figura'),
                       caption='Q-Q plots y distribución de residuos por modelo')
    contenido.append(PageBreak())

    # ── PIE DE PÁGINA
    contenido.append(Spacer(1, 2*cm))
    contenido.append(HRFlowable(width='100%', thickness=1, color=verde))
    contenido.append(Spacer(1, 0.3*cm))
    contenido.append(Paragraph(
        f'Generado automáticamente por el Sistema Inteligente Agrícola — {INSTITUCION} — {FECHA}',
        estilo_nota))

    doc.build(contenido)
    print(f"  [+] PDF generado: {output_path}")
    return str(output_path)


def generar_excel(output_path=None):
    """Genera reporte Excel con múltiples hojas."""
    if output_path is None:
        output_path = OUTPUT_DIR / 'reporte_ia_agricola.xlsx'
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    resultados = cargar_resultados()

    with pd.ExcelWriter(str(output_path), engine='openpyxl') as writer:
        # Hoja 1: Resumen
        resumen_df = pd.DataFrame([{
            'Proyecto': TITULO_PROYECTO,
            'Institución': INSTITUCION,
            'Fecha': FECHA
        }])
        resumen_df.to_excel(writer, sheet_name='Resumen', index=False)

        # Hoja 2: Tabla Comparativa de Modelos
        if resultados.get('modelos') and resultados['modelos'].get('tabla_metricas'):
            df_mod = pd.DataFrame(resultados['modelos']['tabla_metricas'])
            df_mod.to_excel(writer, sheet_name='Comparativa_Modelos', index=False)

        # Hoja 3: Validación Cruzada
        if resultados.get('cv'):
            cv_rows = []
            for nombre_m in ['random_forest', 'xgboost']:
                datos_m = resultados['cv'].get(nombre_m, {})
                if datos_m and datos_m.get('folds'):
                    for fold in datos_m['folds']:
                        fold['modelo'] = nombre_m
                        cv_rows.append(fold)
            if cv_rows:
                pd.DataFrame(cv_rows).to_excel(writer, sheet_name='Cross_Validation', index=False)

        # Hoja 4: Hiperparámetros
        if resultados.get('hiperpar'):
            hp_rows = []
            for m_name in ['random_forest', 'xgboost']:
                datos_hp = resultados['hiperpar'].get(m_name, {})
                if datos_hp and datos_hp.get('best_params'):
                    for k, v in datos_hp['best_params'].items():
                        hp_rows.append({'modelo': m_name, 'hiperparametro': k, 'valor': v})
            if hp_rows:
                pd.DataFrame(hp_rows).to_excel(writer, sheet_name='Hiperparametros', index=False)

        # Hoja 5: Pruebas Estadísticas
        if resultados.get('estadist'):
            est = resultados['estadist']
            for test_name, lista in [
                ('Shapiro_Wilk', est.get('shapiro_wilk', [])),
                ('Ljung_Box',    est.get('ljung_box',    [])),
                ('KS',           est.get('ks',           [])),
                ('Diebold_Mariano', est.get('diebold_mariano', []))
            ]:
                if lista:
                    pd.DataFrame(lista).to_excel(writer, sheet_name=test_name, index=False)

        # Hoja 6: Dataset EDA stats
        if resultados.get('eda') and resultados['eda'].get('estadisticos'):
            try:
                df_eda = pd.DataFrame(resultados['eda']['estadisticos']).T
                df_eda.to_excel(writer, sheet_name='EDA_Estadisticos')
            except Exception:
                pass

    print(f"  [+] Excel generado: {output_path}")
    return str(output_path)


def generar_word(output_path=None):
    """Genera reporte Word con tablas, texto e imágenes."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import docx

    if output_path is None:
        output_path = OUTPUT_DIR / 'reporte_ia_agricola.docx'
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    resultados = cargar_resultados()
    doc = Document()

    # Estilos
    verde_rgb = RGBColor(0x1a, 0x7a, 0x3e)

    def agregar_titulo(texto, nivel=1):
        p = doc.add_heading(texto, level=nivel)
        p.runs[0].font.color.rgb = verde_rgb

    def agregar_parrafo(texto):
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.runs[0].font.size = Pt(10)

    def agregar_figura_word(ruta, caption='', ancho=5.5):
        if ruta and Path(ruta).exists():
            doc.add_picture(str(ruta), width=Inches(ancho))
            if caption:
                p = doc.add_paragraph(f'Figura: {caption}')
                p.runs[0].italic = True
                p.runs[0].font.size = Pt(8)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def agregar_tabla_word(cabecera, datos, col_widths=None):
        tabla = doc.add_table(rows=1 + len(datos), cols=len(cabecera))
        tabla.style = 'Table Grid'
        for j, cab in enumerate(cabecera):
            celda = tabla.rows[0].cells[j]
            celda.text = cab
            celda.paragraphs[0].runs[0].bold = True
        for i, fila in enumerate(datos):
            for j, val in enumerate(fila):
                tabla.rows[i+1].cells[j].text = str(val)
        doc.add_paragraph()

    # ── Portada
    doc.add_heading(TITULO_PROYECTO, 0)
    agregar_parrafo(f'Subtítulo: {SUBTITULO}')
    agregar_parrafo(f'Institución: {INSTITUCION}')
    agregar_parrafo(f'Fecha: {FECHA}')
    doc.add_page_break()

    # ── EDA
    agregar_titulo('1. Análisis Exploratorio de Datos', 1)
    if resultados.get('eda'):
        eda = resultados['eda']
        agregar_parrafo(
            f"Dataset de {eda.get('n_registros', '?')} registros ({eda.get('periodo', '?')}). "
            f"Cultivos analizados: {', '.join(eda.get('cultivos', []))}."
        )
        agregar_figura_word(eda.get('figuras', {}).get('series_temporales'),
                            'Series temporales de precios por cultivo')
        agregar_figura_word(eda.get('figuras', {}).get('heatmap'),
                            'Mapa de calor de correlaciones')
    doc.add_page_break()

    # ── Modelos
    agregar_titulo('2. Comparativa de Modelos', 1)
    if resultados.get('modelos') and resultados['modelos'].get('tabla_metricas'):
        datos_tabla = [
            [f['Modelo'], str(f['RMSE']), str(f['MAE']),
             str(f['R²']), str(f['MAPE %']), f"{f['Tiempo s']}s"]
            for f in resultados['modelos']['tabla_metricas']
        ]
        agregar_tabla_word(
            ['Modelo', 'RMSE', 'MAE', 'R²', 'MAPE %', 'Tiempo'],
            datos_tabla
        )
        agregar_figura_word(resultados['modelos'].get('figuras', {}).get('predicciones_vs_real'),
                            'Predicción vs Real por modelo')
        agregar_figura_word(resultados['modelos'].get('figuras', {}).get('confusion_roc'),
                            'Matriz de Confusión y Curva ROC — Clasificación Plagas')
    doc.add_page_break()

    # ── CV
    agregar_titulo('3. Validación Cruzada', 1)
    if resultados.get('cv'):
        agregar_figura_word(resultados['cv'].get('figura'),
                            f"CV con {resultados['cv'].get('n_folds', 5)} folds")
    doc.add_page_break()

    # ── Pruebas Estadísticas
    agregar_titulo('5. Pruebas Estadísticas', 1)
    if resultados.get('estadist'):
        est = resultados['estadist']
        if est.get('shapiro_wilk'):
            agregar_titulo('Shapiro-Wilk', 2)
            datos_sw = [[r['modelo'], str(r['estadistico']), str(r['p_valor']), r['resultado']]
                        for r in est['shapiro_wilk']]
            agregar_tabla_word(['Modelo', 'Estadístico', 'p-valor', 'Resultado'], datos_sw)

        if est.get('diebold_mariano'):
            agregar_titulo('Diebold-Mariano', 2)
            datos_dm = [[r['modelos'], str(r.get('estadistico_DM', '')),
                         str(r.get('p_valor', '')), r.get('resultado', '')]
                        for r in est['diebold_mariano']]
            agregar_tabla_word(['Modelos', 'DM', 'p-valor', 'Resultado'], datos_dm)

        agregar_figura_word(est.get('figura'), 'Q-Q plots y distribución de residuos')

    doc.save(str(output_path))
    print(f"  [+] Word generado: {output_path}")
    return str(output_path)


def generar_todos(output_dir=None):
    """Genera los tres tipos de reporte."""
    if output_dir:
        output_dir = Path(output_dir)
    else:
        output_dir = OUTPUT_DIR

    print("\nGenerando reportes...")
    rutas = {}
    try:
        rutas['pdf']  = generar_pdf(output_dir / 'reporte_ia_agricola.pdf')
    except Exception as e:
        print(f"  [WARNING] PDF: {e}")
    try:
        rutas['excel'] = generar_excel(output_dir / 'reporte_ia_agricola.xlsx')
    except Exception as e:
        print(f"  [WARNING] Excel: {e}")
    try:
        rutas['word'] = generar_word(output_dir / 'reporte_ia_agricola.docx')
    except Exception as e:
        print(f"  [WARNING] Word: {e}")

    return rutas


if __name__ == '__main__':
    rutas = generar_todos()
    print("\nArchivos generados:")
    for tipo, ruta in rutas.items():
        print(f"  {tipo.upper()}: {ruta}")
